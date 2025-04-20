import asyncio
import os
import json
import re
import hashlib
import time
from typing import List, Dict, Optional, Any, Tuple
from datetime import datetime
from pathlib import Path

# Document processing
from pypdf import PdfReader
from openpyxl import load_workbook
import docx
from dotenv import load_dotenv

# Database and embeddings
import chromadb
from chromadb.utils.embedding_functions import OpenAIEmbeddingFunction

# Caching
from diskcache import Cache as PersistentCache

# LLM
import openai
from openai import OpenAI
import streamlit as st
# Constants
DEBUG = False
DEFAULT_LLM_MODEL = "gpt-4o"  # Default model
CHUNK_SIZE = 1000  # Characters per chunk
CHUNK_OVERLAP = 200  # Overlap between chunks

# =================== LOAD ENVIRONMENT VARIABLES ===================
load_dotenv()
openai_key = os.getenv("OPENAI_API_KEY")
if not openai_key:
    raise ValueError("Please set your OPENAI_API_KEY in the environment variables.")

def sanitize_name(name: str) -> str:
    """
    Sanitize a name to be used as a ChromaDB collection name.
    Rules:
    1. Contains 3-63 characters
    2. Starts and ends with alphanumeric
    3. Contains only alphanumeric, underscores or hyphens
    4. No consecutive periods
    5. Not a valid IPv4 address
    """
    # Replace spaces and other invalid characters with underscores
    sanitized = re.sub(r'[^a-zA-Z0-9-]', '_', name)
    # Remove consecutive underscores
    sanitized = re.sub(r'_+', '_', sanitized)
    # Ensure it starts and ends with alphanumeric
    sanitized = re.sub(r'^[^a-zA-Z0-9]+', '', sanitized)
    sanitized = re.sub(r'[^a-zA-Z0-9]+$', '', sanitized)
    
    # If name is too short, pad it
    if len(sanitized) < 3:
        sanitized = sanitized + "_collection"
    # If name is too long, truncate it
    if len(sanitized) > 63:
        sanitized = sanitized[:63]
        # Ensure it ends with alphanumeric
        sanitized = re.sub(r'[^a-zA-Z0-9]+$', '', sanitized)
    
    return sanitized

# =================== PROJECT RAG CLASS ===================
class ProjectRAG:
    def __init__(self, project_name: str, project_path: str):
        """
        Initialize a RAG system for a specific project
        
        Args:
            project_name: Name of the project (used for the collection name)
            project_path: Path to the project's documents folder
        """
        self.project_name = project_name
        self.project_path = project_path
        self.openai_key = openai_key
        self.client = OpenAI(api_key=self.openai_key)
        self.llm_model_name = os.getenv('LLM_MODEL', DEFAULT_LLM_MODEL)
        
        # Sanitize collection name
        collection_name = sanitize_name(project_name)
        
        # Set up the ChromaDB
        self.db_path = f"./chromadb_storage/{collection_name}"
        os.makedirs(self.db_path, exist_ok=True)
        self.chroma_client = chromadb.PersistentClient(path=self.db_path)
        self.collection = self.chroma_client.get_or_create_collection(
            name=collection_name,
            embedding_function=OpenAIEmbeddingFunction(api_key=openai_key)
        )
        
        # Caching
        self.cache = PersistentCache(cache_dir=f"./cache/{collection_name}", ttl=3600)
        self.response_cache = PersistentCache(cache_dir=f"./response_cache/{collection_name}", ttl=3600)
        
        # Ingestion metadata to avoid reprocessing unchanged files
        self.metadata_path = f"ingestion_metadata_{collection_name}.json"
        self.ingestion_metadata = self.load_ingestion_metadata()
        
        # Statistics
        self.stats = {
            "documents_processed": 0,
            "chunks_stored": 0,
            "last_update": None
        }

    def load_ingestion_metadata(self) -> dict:
        if os.path.exists(self.metadata_path):
            try:
                with open(self.metadata_path, "r", encoding="utf-8") as f:
                    return json.load(f)
            except Exception as e:
                print(f"[ERROR] Loading metadata failed: {e}")
        return {}

    def save_ingestion_metadata(self):
        try:
            with open(self.metadata_path, "w", encoding="utf-8") as f:
                json.dump(self.ingestion_metadata, f)
        except Exception as e:
            print(f"[ERROR] Saving metadata failed: {e}")

    # ------------------ DOCUMENT PREPROCESSING ------------------
    def preprocess_text(self, text: str) -> List[str]:
        """Split text into chunks with overlap"""
        if not text.strip():
            return []
            
        # Simple chunking by characters with overlap
        chunks = []
        for i in range(0, len(text), CHUNK_SIZE - CHUNK_OVERLAP):
            chunk = text[i:i + CHUNK_SIZE]
            if chunk.strip():
                chunks.append(chunk)
        return chunks

    # ------------------ DOCUMENT INGESTION METHODS ------------------
    async def extract_text_from_pdf(self, pdf_path: str) -> str:
        """Extract all text from a PDF file"""
        try:
            reader = PdfReader(pdf_path)
            text = ""
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n\n"
            return text
        except Exception as e:
            print(f"[ERROR] Failed to extract PDF {pdf_path}: {e}")
            return ""

    async def extract_text_from_docx(self, docx_path: str) -> str:
        """Extract all text from a Word document"""
        try:
            doc = docx.Document(docx_path)
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text)
                    text += " | ".join(row_text) + "\n"
            return text
        except Exception as e:
            print(f"[ERROR] Failed to extract DOCX {docx_path}: {e}")
            return ""

    async def extract_data_from_excel(self, excel_path: str) -> Tuple[str, List[str]]:
        """
        Extract all data from Excel as text, including sheet names and file path context
        Returns tuple of (text content, list of sheet names)
        """
        try:
            print(f"[INFO] Processing Excel file: {excel_path}")
            # Add timeout protection
            wb = load_workbook(excel_path, data_only=True, read_only=True)
            
            text = []
            sheet_names = []
            
            # Add file path context
            file_name = os.path.basename(excel_path)
            parent_folder = os.path.basename(os.path.dirname(excel_path))
            text.append(f"File: {file_name}")
            text.append(f"Location: {parent_folder}")
            
            for sheet_name in wb.sheetnames:
                sheet_names.append(sheet_name)
                sheet = wb[sheet_name]
                text.append(f"\nSheet: {sheet_name}")
                
                # Process rows with a limit to prevent hanging
                row_count = 0
                max_rows = 1000  # Limit number of rows to process
                
                for row in sheet.iter_rows(values_only=True):
                    row_count += 1
                    if row_count > max_rows:
                        text.append(f"[Note: Truncated after {max_rows} rows]")
                        break
                        
                    # Filter out empty cells and convert to strings
                    row_values = [str(cell) if cell is not None else "" for cell in row]
                    # Only add rows that have some content
                    if any(val.strip() for val in row_values):
                        text.append(" | ".join(row_values))
            
            wb.close()
            return "\n".join(text), sheet_names
            
        except Exception as e:
            print(f"[ERROR] Failed to extract Excel {excel_path}: {e}")
            return "", []

    async def ingest_document(self, file_path: str) -> bool:
        """Ingest a document with enhanced context from file path"""
        try:
            # Check if file has been modified since last ingestion
            mod_time = os.path.getmtime(file_path)
            if file_path in self.ingestion_metadata and self.ingestion_metadata[file_path] >= mod_time:
                print(f"[INFO] File unchanged: {file_path}")
                return False
                
            # Get file context
            file_name = os.path.basename(file_path)
            parent_folder = os.path.basename(os.path.dirname(file_path))
            relative_path = os.path.relpath(file_path, self.project_path)
            
            # Extract text based on file type
            ext = os.path.splitext(file_path)[1].lower()
            document_text = ""
            metadata = {
                "source": file_path,
                "file_name": file_name,
                "parent_folder": parent_folder,
                "relative_path": relative_path,
                "file_type": ext.replace(".", ""),
                "timestamp": datetime.now().isoformat()
            }
            
            if ext == ".pdf":
                document_text = await self.extract_text_from_pdf(file_path)
                document_text = f"File: {file_name}\nLocation: {parent_folder}\n\n{document_text}"
                
            elif ext in [".docx", ".doc"]:
                document_text = await self.extract_text_from_docx(file_path)
                document_text = f"File: {file_name}\nLocation: {parent_folder}\n\n{document_text}"
                
            elif ext in [".xlsx", ".xls"]:
                document_text, sheet_names = await self.extract_data_from_excel(file_path)
                metadata["sheet_names"] = sheet_names
                
            elif ext == ".txt":
                with open(file_path, "r", encoding="utf-8") as f:
                    document_text = f.read()
                document_text = f"File: {file_name}\nLocation: {parent_folder}\n\n{document_text}"
            else:
                print(f"[WARN] Unsupported file type: {file_path}")
                return False
                
            if not document_text.strip():
                print(f"[WARN] No content extracted from: {file_path}")
                return False
                
            # Process the text into chunks
            chunks = self.preprocess_text(document_text)
            if not chunks:
                print(f"[WARN] No chunks created for: {file_path}")
                return False
                
            # Add chunks to the database
            for i, chunk in enumerate(chunks):
                chunk_id = f"{sanitize_name(file_name)}_{i}"
                chunk_metadata = metadata.copy()
                chunk_metadata["chunk_index"] = i
                chunk_metadata["total_chunks"] = len(chunks)
                
                # First try to delete existing chunk if it exists
                try:
                    self.collection.delete(ids=[chunk_id])
                except:
                    pass
                    
                # Add the new chunk
                self.collection.add(
                    ids=[chunk_id],
                    documents=[chunk],
                    metadatas=[chunk_metadata]
                )
                
            # Update metadata and stats
            self.ingestion_metadata[file_path] = mod_time
            self.save_ingestion_metadata()
            self.stats["documents_processed"] += 1
            self.stats["chunks_stored"] += len(chunks)
            self.stats["last_update"] = datetime.now().isoformat()
            
            print(f"[INFO] Successfully ingested {file_path} with {len(chunks)} chunks")
            return True
            
        except Exception as e:
            print(f"[ERROR] Failed to ingest {file_path}: {e}")
            return False

    async def ingest_directory(self) -> Dict[str, Any]:
        """Ingest all supported documents in the project directory"""
        start_time = time.time()
        processed_count = 0
        skipped_count = 0
        error_count = 0
        
        print(f"[INFO] Starting ingestion for project: {self.project_name}")
        print(f"[INFO] Scanning directory: {self.project_path}")
        
        # Track ingestion metrics
        ingestion_results = {
            "processed_files": [],
            "skipped_files": [],
            "error_files": [],
            "start_time": datetime.now().isoformat()
        }
        
        # Walk through all files in the directory and its subdirectories
        for root, _, files in os.walk(self.project_path):
            for file in files:
                file_path = os.path.join(root, file)
                try:
                    # Check if file extension is supported
                    ext = os.path.splitext(file_path)[1].lower()
                    if ext not in [".pdf", ".docx", ".doc", ".xlsx", ".xls", ".txt"]:
                        continue
                        
                    # Get relative path from project root for metadata
                    rel_path = os.path.relpath(file_path, self.project_path)
                    
                    success = await self.ingest_document(file_path)
                    if success:
                        processed_count += 1
                        ingestion_results["processed_files"].append({
                            "file": rel_path,
                            "full_path": file_path
                        })
                    else:
                        skipped_count += 1
                        ingestion_results["skipped_files"].append({
                            "file": rel_path,
                            "full_path": file_path
                        })
                except Exception as e:
                    error_count += 1
                    ingestion_results["error_files"].append({
                        "file": rel_path,
                        "full_path": file_path,
                        "error": str(e)
                    })
                    print(f"[ERROR] Failed to process {file_path}: {e}")
        
        # Calculate elapsed time
        elapsed_time = time.time() - start_time
        ingestion_results["elapsed_time"] = elapsed_time
        ingestion_results["end_time"] = datetime.now().isoformat()
        ingestion_results["total_processed"] = processed_count
        ingestion_results["total_skipped"] = skipped_count
        ingestion_results["total_errors"] = error_count
        
        print(f"[INFO] Ingestion completed for {self.project_name}")
        print(f"[INFO] Processed: {processed_count}, Skipped: {skipped_count}, Errors: {error_count}")
        print(f"[INFO] Elapsed time: {elapsed_time:.2f} seconds")
        
        return ingestion_results

    # ------------------ QUERY & RESPONSE METHODS ------------------
    async def query_collection(self, query: str, n_results: int = 5) -> List[Dict[str, Any]]:
        """
        Query the collection and return the most relevant chunks with metadata
        """
        try:
            # Generate cache key
            query_hash = hashlib.md5(query.encode()).hexdigest()
            cached = self.cache.get(query_hash)
            if cached:
                print(f"[INFO] cached: {cached}")
                if DEBUG:
                    print(f"[DEBUG] Using cached chunks for query: {query}")
                return cached
                
            # Query the collection
            results = self.collection.query(
                query_texts=[query], 
                n_results=n_results, 
                include=["documents", "metadatas", "distances"]
            )
            print(f"[INFO] results in query_collection: {results}")

            if not results["documents"] or not results["documents"][0]:
                return []
                
            # Format the results with metadata
            retrieved = []
            for i, doc in enumerate(results["documents"][0]):
                item = {
                    "content": doc,
                    "metadata": results["metadatas"][0][i] if results["metadatas"] else {},
                    "relevance_score": results["distances"][0][i] if results["distances"] else None
                }
                retrieved.append(item)
                
            # Cache the results
            self.cache.set(query_hash, retrieved)
            
            if DEBUG:
                print(f"[DEBUG] Found {len(retrieved)} chunks for query: {query}")
            return retrieved
        except Exception as e:
            print(f"[ERROR] Error retrieving data for '{query}': {e}")
            return []

    async def generate_response(self, query: str, context_chunks: List[Dict[str, Any]]) -> Dict[str, Any]:
        """
        Generate a response based on the query and retrieved context chunks
        """
        # Generate cache key
        query_hash = hashlib.md5(query.encode()).hexdigest()
        cached_response = self.response_cache.get(query_hash)
        if cached_response:
            if DEBUG:
                print(f"[DEBUG] Using cached response for query: {query}")
            return cached_response
            
        # Format context for the prompt
        formatted_context = ""
        sources = []
        chunks = ""
        print(f"[INFO] context_chunks: {len(context_chunks)}")
        for i, chunk in enumerate(context_chunks):
            chunks += str(chunk["metadata"]["chunk_index"]) + ", "
            
            formatted_context += f"[CHUNK {i+1}] {chunk['content']}\n\n"
            if "metadata" in chunk and "source" in chunk["metadata"]:
                source_file = os.path.basename(chunk["metadata"]["source"])
                if source_file not in sources:
                    sources.append(source_file)
                    
        if not formatted_context:
            formatted_context = "No relevant information found in the project documents."
        print(f"[INFO] formatted_context: {formatted_context}")
        # Create prompt for the LLM
        system_prompt = (
            "You are an AI assistant specialized in analyzing grant applications and project documents. "
            "You will be provided with context chunks from a project's documents. "
            "Use this information to answer the query accurately and concisely. "
            "If the information is not in the context, state that clearly. "
            "Include specific facts, figures, and quotes from the documents when relevant. "
            "Always cite your sources when quoting from specific documents."
        )
        
        user_prompt = (
            f"Query: {query}\n\n"
            f"Context from project documents:\n{formatted_context}"
        )
        
        try:
            # Call the LLM
            response = self.client.chat.completions.create(
                model=self.llm_model_name,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                temperature=0.3
            )
            
            # Extract and cache the response
            answer = response.choices[0].message.content
            
            result = {
                "answer": answer,
                "sources": sources,
                "timestamp": datetime.now().isoformat(),
                "context_used": len(context_chunks),
                "chunks": chunks
            }
            
            self.response_cache.set(query_hash, result)
            return result 
            
        except Exception as e:
            print(f"[ERROR] Failed to generate response: {e}")
            return {
                "answer": f"Error generating response: {str(e)}",
                "sources": [],
                "timestamp": datetime.now().isoformat(),
                "error": str(e)
            }

    async def ask(self, query: str) -> Dict[str, Any]:
        """
        Main method to ask a question about the project
        """
        print(f"[INFO] Processing query for {self.project_name}: {query}")
        
        # 1. Retrieve relevant chunks
        retrieved_chunks = await self.query_collection(query, n_results=5)
    
        response = await self.generate_response(query, retrieved_chunks)
        
        return response

    # ------------------ REPORT GENERATION METHODS ------------------
    async def check_eligibility(self, criteria: Dict[str, str]) -> Dict[str, Any]:
        """
        Check project's eligibility based on specified criteria
        
        Args:
            criteria: Dictionary mapping criteria names to questions
            
        Returns:
            Dictionary with criteria names, questions, answers, and eligibility assessment
        """
        results = {
            "project_name": self.project_name,
            "timestamp": datetime.now().isoformat(),
            "criteria": [],
            "eligible": True,  # Initially assume eligible
            "summary": ""
        }
        
        for criterion_name, question in criteria.items():
            print(f"[INFO] Checking criterion '{criterion_name}' for {self.project_name}")
            
            # Format the question to explicitly ask about eligibility
            eligibility_question = (
                f"Based on the project documents, {question} "
                f"Answer with 'Yes' or 'No' first, then provide supporting evidence."
            )
            
            # Get answer from the RAG system
            response = await self.ask(eligibility_question)
            
            # Determine eligibility by checking if the answer starts with "Yes"
            answer = response["answer"].strip()
            is_eligible = answer.lower().startswith("yes")
            
            # If any criterion fails, the project is not eligible
            if not is_eligible:
                results["eligible"] = False
                
            # Add criterion result
            results["criteria"].append({
                "name": criterion_name,
                "question": question,
                "answer": answer,
                "meets_criterion": is_eligible,
                "sources": response.get("sources", [])
            })
        
        # Generate summary
        if results["eligible"]:
            results["summary"] = f"Project '{self.project_name}' meets all eligibility criteria."
        else:
            failed_criteria = [c["name"] for c in results["criteria"] if not c["meets_criterion"]]
            results["summary"] = (
                f"Project '{self.project_name}' does not meet the following criteria: "
                f"{', '.join(failed_criteria)}."
            )
            
        return results

    async def check_selected_projects(self, criteria: Dict[str, str]) -> Dict[str, Any]:
        """
        Select projects based on specified criteria and return evaluation with actions needed.

        Args:
            criteria: Dictionary mapping criteria names to questions

        Returns:
            Dictionary with criteria names, questions, answers, selection status,
            and actions needed to meet criteria
        """
        results = {
            "project_name": self.project_name,
            "timestamp": datetime.now().isoformat(),
            "criteria": [],
            "selected": True,  # Assume selected unless a criterion fails
            "summary": ""
        }

        for criterion_name, question in criteria.items():
            print(f"[INFO] Checking criterion '{criterion_name}' for {self.project_name}")

            # Format selection question
            selection_question = (
                f"Based on the project documents, {question} "
                f"Answer with 'Yes' or 'No' first, then provide supporting evidence."
            )

            response = await self.ask(selection_question)
            answer = response["answer"].strip()
            is_selected = answer.lower().startswith("yes")

            # If any criterion fails, the project is not selected
            if not is_selected:
                results["selected"] = False

            # Get action needed if the criterion is not met
            if not is_selected:
                action_prompt = (
                    f"The project does not meet the following criterion: '{question}'. "
                    f"What specific actions should be taken to meet this requirement?"
                )
                action_response = await self.ask(action_prompt)
                action_needed = action_response["answer"].strip()
            else:
                action_needed = "No action needed."

            # Add full criterion result
            results["criteria"].append({
                "name": criterion_name,
                "question": question,
                "answer": answer,
                "meets_criterion": is_selected,
                "sources": response.get("sources", []),
                "action_needed": action_needed
            })

        # Summary
        if results["selected"]:
            results["summary"] = f"Project '{self.project_name}' meets all selection criteria."
        else:
            failed_criteria = [c["name"] for c in results["criteria"] if not c["meets_criterion"]]
            results["summary"] = (
                f"Project '{self.project_name}' does not meet the following criteria: "
                f"{', '.join(failed_criteria)}."
            )

        return results
    
    async def generate_detailed_report(self, questions: List[str]) -> Dict[str, Any]:
        """
        Generate a detailed report by answering a list of questions about the project
        
        Args:
            questions: List of questions to answer about the project
            
        Returns:
            Dictionary with project name, timestamp, and answers to questions
        """
        report = {
            "project_name": self.project_name,
            "timestamp": datetime.now().isoformat(),
            "sections": []
        }
        
        for question in questions:
            print(f"[INFO] Answering report question for {self.project_name}: {question}")
            
            response = await self.ask(question)
            
            report["sections"].append({
                "question": question,
                "answer": response["answer"],
                "sources": response.get("sources", [])
            })
            
        return report

    async def generate_recommendation(self, eligibility_result: Dict[str, Any], detailed_report: Dict[str, Any]) -> Dict[str, Any]:
        """
        Generate a recommendation for the donor based on eligibility and detailed report
        
        Args:
            eligibility_result: Result from check_eligibility
            detailed_report: Result from generate_detailed_report
            
        Returns:
            Dictionary with recommendation details
        """
        # Prepare context about the project for the LLM
        context = f"Project Name: {self.project_name}\n\n"
        
        # Add eligibility information
        context += "ELIGIBILITY ASSESSMENT:\n"
        context += f"Overall Eligibility: {eligibility_result['eligible']}\n"
        for criterion in eligibility_result["criteria"]:
            context += f"- {criterion['name']}: {'Meets criterion' if criterion['meets_criterion'] else 'Does not meet criterion'}\n"
            context += f"  Question: {criterion['question']}\n"
            context += f"  Answer: {criterion['answer']}\n\n"
            
        # Add report information
        context += "DETAILED REPORT:\n"
        for section in detailed_report["sections"]:
            context += f"Question: {section['question']}\n"
            context += f"Answer: {section['answer']}\n\n"
            
        # Prepare prompt for recommendation generation
        system_prompt = (
            "You are a grant evaluation expert assisting a donor in making funding decisions. "
            "Your role is to provide an objective recommendation based on project eligibility and "
            "detailed assessment, highlighting strengths, weaknesses, risks, and potential impact. "
            "Your recommendation should be clear, substantiated with evidence from the project documents, "
            "and include specific funding suggestions or alternatives if appropriate. "
            "You MUST start your response with one of these exact phrases on the first line:\n"
            "DECISION: Fund\n"
            "DECISION: Do Not Fund\n"
            "DECISION: Partially Fund\n"
        )
        
        user_prompt = (
            "Based on the following project assessment, provide a donor recommendation that includes:\n"
            "1. Funding decision (Must start with DECISION: followed by Fund/Do Not Fund/Partially Fund)\n"
            "2. Executive summary (2-3 sentences)\n"
            "3. Key strengths and weaknesses\n"
            "4. Risks and mitigations\n"
            "5. Expected impact if funded\n"
            "6. Any conditions or special considerations\n\n"
            f"{context}"
        )
        
        try:
            # Generate recommendation using OpenAI
            response = self.client.chat.completions.create(
                model=self.llm_model_name,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                temperature=0.2
            )
            
            recommendation_text = response.choices[0].message.content
            
            # Extract the funding decision from the first line
            first_line = recommendation_text.split('\n')[0].strip()
            funding_decision = "Pending"
            if first_line.startswith("DECISION:"):
                funding_decision = first_line.replace("DECISION:", "").strip()
                # Remove the decision line from the recommendation text
                recommendation_text = '\n'.join(recommendation_text.split('\n')[1:]).strip()
            
            recommendation = {
                "project_name": self.project_name,
                "timestamp": datetime.now().isoformat(),
                "funding_decision": funding_decision,
                "recommendation": recommendation_text
            }
            
            return recommendation
            
        except Exception as e:
            print(f"[ERROR] Failed to generate recommendation: {e}")
            return {
                "project_name": self.project_name,
                "timestamp": datetime.now().isoformat(),
                "error": str(e),
                "recommendation": "Error generating recommendation."
            }

# =================== GRANT ASSESSMENT SYSTEM ===================
class GrantAssessmentSystem:
    def __init__(self, projects_dir: str):
        """
        Initialize the grant assessment system
        
        Args:
            projects_dir: Directory containing project folders
        """
        self.projects_dir = projects_dir
        self.projects = {}  # Map of project_name -> ProjectRAG
        self.openai_key = openai_key
        self.client = OpenAI(api_key=self.openai_key)
        self.llm_model_name = os.getenv('LLM_MODEL', DEFAULT_LLM_MODEL)
        
        # Create projects directory if it doesn't exist
        os.makedirs(projects_dir, exist_ok=True)

    async def initialize_projects(self):
        """Initialize RAG systems for all projects in the projects directory"""
        try:
            # Scan projects directory
            if not os.path.exists(self.projects_dir):
                print(f"[INFO] Creating projects directory: {self.projects_dir}")
                os.makedirs(self.projects_dir)
                return

            # Initialize ProjectRAG for each project folder
            for item in os.listdir(self.projects_dir):
                project_path = os.path.join(self.projects_dir, item)
                if os.path.isdir(project_path):
                    print(f"[INFO] Initializing project: {item}")
                    self.projects[item] = ProjectRAG(item, project_path)

            print(f"[INFO] Initialized {len(self.projects)} projects")
            
        except Exception as e:
            print(f"[ERROR] Failed to initialize projects: {e}")
            raise

    async def ingest_project(self, project_name: str) -> bool:
        """
        Ingest all documents for a specific project
        
        Args:
            project_name: Name of the project to ingest
            
        Returns:
            bool: True if successful, False otherwise
        """
        if project_name not in self.projects:
            print(f"[ERROR] Project not found: {project_name}")
            return False
            
        try:
            # Reset project stats before ingestion
            self.projects[project_name].stats = {
                "documents_processed": 0,
                "chunks_stored": 0,
                "last_update": None
            }
            
            # Start ingestion
            start_time = time.time()
            results = await self.projects[project_name].ingest_directory()
            
            # Update project stats
            project = self.projects[project_name]
            project.stats["last_update"] = datetime.now().isoformat()
            
            # Calculate processing metrics
            elapsed_time = time.time() - start_time
            avg_time_per_doc = elapsed_time / max(1, project.stats["documents_processed"])
            
            # Store metrics in session state if available
            try:
                import streamlit as st
                if "processing_metrics" in st.session_state:
                    st.session_state.processing_metrics[project_name] = {
                        "Documents Processed": project.stats["documents_processed"],
                        "Chunks Stored": project.stats["chunks_stored"],
                        "Processing Time": f"{elapsed_time:.1f}s",
                        "Average Time per Document": f"{avg_time_per_doc:.2f}s"
                    }
                if "operation_timestamps" in st.session_state:
                    if project_name not in st.session_state.operation_timestamps:
                        st.session_state.operation_timestamps[project_name] = {}
                    st.session_state.operation_timestamps[project_name]["Last Ingestion"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                if "ingested_projects" in st.session_state:
                    st.session_state.ingested_projects.add(project_name)
            except:
                pass  # Streamlit context may not be available
                
            print(f"[INFO] Successfully ingested project {project_name}")
            print(f"[INFO] Documents processed: {project.stats['documents_processed']}")
            print(f"[INFO] Chunks stored: {project.stats['chunks_stored']}")
            print(f"[INFO] Processing time: {elapsed_time:.1f}s")
            
            return True
            
        except Exception as e:
            print(f"[ERROR] Failed to ingest project {project_name}: {e}")
            return False

    async def ingest_all_projects(self) -> Dict[str, Any]:
        """
        Ingest all documents for all projects
        
        Returns:
            Dictionary with ingestion results for each project
        """
        results = {}
        for project_name in self.projects:
            print(f"[INFO] Ingesting project: {project_name}")
            success = await self.ingest_project(project_name)
            results[project_name] = {"success": success}
        return results

    async def add_project_folder(self, folder_path: str) -> bool:
        """
        Add a new project folder to the system
        
        Args:
            folder_path: Path to the project folder to add
            
        Returns:
            bool: True if successfully added, False otherwise
        """
        try:
            if not os.path.isdir(folder_path):
                print(f"[ERROR] Not a valid directory: {folder_path}")
                return False
                
            project_name = os.path.basename(folder_path)
            if project_name in self.projects:
                print(f"[WARN] Project {project_name} already exists")
                return False
                
            # Copy folder to projects directory
            target_path = os.path.join(self.projects_dir, project_name)
            if os.path.exists(target_path):
                print(f"[WARN] Target path already exists: {target_path}")
                return False
                
            # Create target directory and copy files
            os.makedirs(target_path, exist_ok=True)
            for root, _, files in os.walk(folder_path):
                for file in files:
                    src_file = os.path.join(root, file)
                    rel_path = os.path.relpath(src_file, folder_path)
                    dst_file = os.path.join(target_path, rel_path)
                    os.makedirs(os.path.dirname(dst_file), exist_ok=True)
                    import shutil
                    shutil.copy2(src_file, dst_file)
                    
            # Initialize ProjectRAG for the new folder
            self.projects[project_name] = ProjectRAG(project_name, target_path)
            print(f"[INFO] Successfully added project: {project_name}")
            
            # Ingest the new project
            await self.ingest_project(project_name)
            return True
            
        except Exception as e:
            print(f"[ERROR] Failed to add project folder: {e}")
            return False

    async def chat_with_projects(self, query: str, project_names: list[str]) -> dict:
        """
        Ask a question to multiple projects and generate a comparative analysis.
        
        Args:
            query (str): The question to ask each project
            project_names (list[str]): List of project names to query
            
        Returns:
            dict: Contains individual project responses and comparative analysis
        """
        if not project_names or len(project_names) < 2:
            raise ValueError("At least 2 projects are required for multi-project chat")
        
        # Get responses from each project
        responses = {}
        for project_name in project_names:
            try:
                project_response = await self.ask_project(project_name, query)
                responses[project_name] = project_response
            except Exception as e:
                responses[project_name] = {
                    "answer": f"Error getting response: {str(e)}",
                    "sources": []
                }
        
        # Generate comparative analysis
        try:
            # Prepare context for comparison
            context = f"Question asked: {query}\n\nProject responses:\n"
            for project, response in responses.items():
                context += f"\n{project}:\n{response['answer']}"
            
            comparison_prompt = f"""Based on the responses from multiple projects to the question "{query}", please provide a comparative analysis.
            Focus on:
            1. Key similarities and differences in the responses
            2. Notable insights unique to each project
            3. Overall patterns or trends
            4. Implications of these differences
            
            Context:
            {context}
            
            Please provide a clear, structured analysis that helps understand how the projects relate to each other in the context of this question."""
            
            response = await self.client.chat.completions.create(
                model=self.llm_model_name,
                messages=[
                    {"role": "system", "content": "You are an expert grant analyst tasked with comparing responses from multiple projects. Provide clear, insightful analysis that helps understand the relationships and differences between projects."},
                    {"role": "user", "content": comparison_prompt}
                ],
                temperature=0.3
            )
            
            comparison = response.choices[0].message.content
            
        except Exception as e:
            comparison = f"Error generating comparative analysis: {str(e)}"
        
        return {
            "responses": responses,
            "comparison": comparison
        }

    async def search_projects(self, query: str) -> Dict[str, List[Dict[str, Any]]]:
        """
        Search across all projects for relevant information
        
        Args:
            query: Search query
            
        Returns:
            Dictionary mapping project names to lists of relevant chunks
        """
        results = {}
        for project_name, project in self.projects.items():
            chunks = await project.query_collection(query, n_results=3)
            if chunks:
                results[project_name] = chunks
                
        return results

    async def generate_comparative_analysis(self, eligible_only: bool = True) -> Dict[str, Any]:
        """
        Generate a comparative analysis of multiple projects
        
        Args:
            eligible_only: If True, only compare eligible projects
            
        Returns:
            Dictionary containing comparative analysis results
        """
        try:
            if len(self.projects) < 2:
                return {
                    "error": "At least two projects are required for comparative analysis",
                    "timestamp": datetime.now().isoformat()
                }
            
            # Prepare context about all projects
            projects_context = ""
            responses = {}
            
            for project_name, project in self.projects.items():
                # Query each project for key information
                query = (
                    "Summarize this project's key aspects including: "
                    "1. Main objectives and goals "
                    "2. Target beneficiaries "
                    "3. Implementation approach "
                    "4. Expected outcomes and impact "
                    "5. Budget and resource requirements"
                )
                
                response = await project.ask(query)
                responses[project_name] = response
                
                projects_context += f"\nProject: {project_name}\n{response['answer']}\n"
                
            # Generate comparative analysis using GPT-4
            system_prompt = (
                "You are an expert grant analyst tasked with comparing multiple projects. "
                "Provide a detailed comparative analysis focusing on strengths, weaknesses, "
                "synergies, and potential impact. Be objective and support your analysis "
                "with specific examples from the projects."
            )
            
            user_prompt = (
                "Compare the following projects, analyzing their relative merits, "
                "potential impact, and areas of complementarity or overlap:\n\n"
                f"{projects_context}\n\n"
                "Please structure your analysis to cover:\n"
                "1. Key similarities and differences\n"
                "2. Relative strengths and weaknesses\n"
                "3. Potential synergies or overlaps\n"
                "4. Comparative impact assessment\n"
                "5. Resource efficiency comparison\n"
                "6. Recommendations for optimization"
            )
            
            # Generate comparative analysis
            analysis_response = self.client.chat.completions.create(
                model=self.llm_model_name,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                temperature=0.3
            )
            
            return {
                "responses": responses,
                "comparison": analysis_response.choices[0].message.content,
                "timestamp": datetime.now().isoformat(),
                "projects_compared": list(self.projects.keys())
            }
            
        except Exception as e:
            print(f"[ERROR] Failed to generate comparative analysis: {e}")
            return {
                "error": str(e),
                "timestamp": datetime.now().isoformat()
            }

    async def ask_project(self, project_name: str, question: str) -> Dict[str, Any]:
        """
        Ask a question to a specific project
        
        Args:
            project_name: Name of the project to query
            question: Question to ask
            
        Returns:
            Dictionary with the response from the project
        """
        if project_name not in self.projects:
            return {
                "error": f"Project {project_name} not found",
                "answer": f"Error: Project {project_name} not found",
                "sources": [],
                "timestamp": datetime.now().isoformat()
            }
            
        try:
            response = await self.projects[project_name].ask(question)
            return response
        except Exception as e:
            print(f"[ERROR] Failed to query project {project_name}: {e}")
            return {
                "error": str(e),
                "answer": f"Error querying project: {str(e)}",
                "sources": [],
                "timestamp": datetime.now().isoformat()
            }

# =================== MAIN FUNCTION ===================
async def main():
    # Setup the grant assessment system
    system = GrantAssessmentSystem("./GrantRAG/projects_data")
    await system.initialize_projects()
    
    # Ingest all project documents
    print("[INFO] Starting document ingestion for all projects...")
    await system.ingest_all_projects()
    print("[INFO] Document ingestion completed")
    
    # Example of checking eligibility for all projects
    print("[INFO] Checking eligibility for all projects...")
    eligibility_results = await system.check_all_projects_eligibility()
    for project_name, result in eligibility_results.items():
        print(f"Project '{project_name}' eligible: {result['eligible']}")
    
    # Example of generating recommendations
    print("[INFO] Generating recommendations for all projects...")
    recommendations = await system.generate_all_recommendations()
    
    # Example of comparative analysis
    print("[INFO] Generating comparative analysis...")
    analysis = await system.generate_comparative_analysis()
    print("[INFO] Analysis completed")
   
if __name__ == "__main__":
    asyncio.run(main()) 